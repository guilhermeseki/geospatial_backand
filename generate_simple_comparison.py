#!/usr/bin/env python3
"""
Comparação Simples: On-Premise vs Nuvem
Apenas vantagens e desvantagens
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()

def create_simple_comparison():
    """Create simplified comparison document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ======================
    # TITLE
    # ======================
    title = doc.add_heading('Comparação: Comprar Hardware vs Nuvem', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.add_run('Análise de Vantagens e Desvantagens').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n')
    info.add_run('Sistema: API Geoespacial (FastAPI + GeoServer + Dask)\n')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('_' * 80)

    add_page_break(doc)

    # ======================
    # RESUMO DE CUSTOS
    # ======================
    doc.add_heading('💰 Resumo de Custos', 1)

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header
    hdr_cells = table.rows[0].cells
    headers = ['Opção', 'Investimento Inicial', 'Custo Mensal', 'Total 5 Anos']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_background(hdr_cells[i], '4472C4')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    data = [
        ('On-Premise (Hardware Novo)', 'R$ 8.800', 'R$ 693', 'R$ 53.380'),
        ('Oracle Cloud (Backfill)', 'R$ 0', 'R$ 482', 'R$ 29.000'),
        ('Híbrido Ultra-Otimizado', 'R$ 0', 'R$ 143', 'R$ 8.580'),
    ]

    for i, row_data in enumerate(data):
        row_cells = table.add_row().cells
        for j, value in enumerate(row_data):
            row_cells[j].text = value
            # Highlight best option
            if i == 2:  # Híbrido
                set_cell_background(row_cells[j], 'FFD700')

    doc.add_paragraph()

    # Key findings
    doc.add_heading('Principais Conclusões:', 2)
    key = doc.add_paragraph(style='List Bullet')
    key.add_run('💸 Investimento inicial: Nuvem R$ 0 vs Hardware R$ 8.800\n')
    key.add_run('📊 Custo mensal: Híbrido R$ 143 vs On-premise R$ 693\n')
    key.add_run('📈 5 anos: Híbrido R$ 8.580 vs On-premise R$ 53.380\n')
    key.add_run('🎯 Economia máxima: R$ 44.800 em 5 anos (84%)\n')

    add_page_break(doc)

    # ======================
    # ON-PREMISE
    # ======================
    doc.add_heading('🖥️ Opção 1: Comprar Hardware (On-Premise)', 1)

    # Specs
    doc.add_heading('Configuração Recomendada:', 2)
    specs = doc.add_paragraph(style='List Bullet')
    specs.add_run('Processador: Intel i7-13700 ou AMD Ryzen 7 5800X\n')
    specs.add_run('RAM: 64 GB DDR4 (2×32GB)\n')
    specs.add_run('SSD NVMe: 1TB (sistema + cache)\n')
    specs.add_run('HDD: 2× 6TB (dados + backup)\n')
    specs.add_run('Fonte: 650W 80+ Gold\n')
    specs.add_run('UPS: 1500VA\n')

    doc.add_paragraph()

    # Cost breakdown
    doc.add_heading('Custos:', 2)

    doc.add_paragraph().add_run('Investimento Inicial:').bold = True
    initial = doc.add_paragraph(style='List Bullet')
    initial.add_run('Hardware completo: R$ 8.800\n')

    doc.add_paragraph().add_run('Custos Mensais:').bold = True
    monthly = doc.add_paragraph(style='List Bullet')
    monthly.add_run('Energia (400W, 24/7): R$ 245/mês\n')
    monthly.add_run('Internet business: R$ 300/mês\n')
    monthly.add_run('Manutenção (10%/ano): R$ 73/mês\n')
    monthly.add_run('Backup externo: R$ 75/mês\n')
    monthly.add_run('TOTAL MENSAL: R$ 693/mês\n').bold = True

    doc.add_paragraph().add_run('Total 5 Anos: R$ 53.380').bold = True

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    # VANTAGENS
    doc.add_heading('✅ VANTAGENS', 2)
    vant_onprem = doc.add_paragraph(style='List Bullet')
    vant_onprem.add_run('Controle total sobre hardware e software\n')
    vant_onprem.add_run('Latência zero (tudo local)\n')
    vant_onprem.add_run('Dados ficam 100% sob seu controle físico\n')
    vant_onprem.add_run('Sem dependência de internet para acessar dados\n')
    vant_onprem.add_run('Sem vendor lock-in (não fica preso a um provedor)\n')
    vant_onprem.add_run('Custo mensal previsível (R$ 693)\n')
    vant_onprem.add_run('Pode usar 100% dos recursos quando quiser\n')
    vant_onprem.add_run('Compliance LGPD simplificado (dados no Brasil)\n')

    # DESVANTAGENS
    doc.add_heading('❌ DESVANTAGENS', 2)
    desv_onprem = doc.add_paragraph(style='List Bullet')
    desv_onprem.add_run('Investimento inicial alto (R$ 8.800)\n')
    desv_onprem.add_run('Você é responsável por toda manutenção\n')
    desv_onprem.add_run('Hardware pode falhar (risco de perda de dados)\n')
    desv_onprem.add_run('Backup manual ou semi-automático\n')
    desv_onprem.add_run('Sem SLA garantido (disponibilidade ~95%)\n')
    desv_onprem.add_run('Escalabilidade limitada (precisa comprar mais hardware)\n')
    desv_onprem.add_run('Hardware deprecia (vida útil 3-5 anos)\n')
    desv_onprem.add_run('Precisa de espaço físico e refrigeração\n')
    desv_onprem.add_run('Custo de energia (R$ 245/mês = R$ 2.940/ano)\n')
    desv_onprem.add_run('Sem disaster recovery automático\n')
    desv_onprem.add_run('Acesso remoto requer configuração (VPN, DNS dinâmico)\n')
    desv_onprem.add_run('Upgrades futuros custam mais (RAM, HDD, etc)\n')

    add_page_break(doc)

    # ======================
    # NUVEM - ORACLE BACKFILL
    # ======================
    doc.add_heading('☁️ Opção 2: Oracle Cloud (Estratégia Backfill)', 1)

    # Specs
    doc.add_heading('Configuração:', 2)
    specs_cloud = doc.add_paragraph(style='List Bullet')
    specs_cloud.add_run('Servidor 24/7: 4 OCPU, 16 GB RAM (serving)\n')
    specs_cloud.add_run('Worker on-demand: 16 OCPU, 64 GB RAM (processing)\n')
    specs_cloud.add_run('Storage: 2.5 TB Object Storage\n')
    specs_cloud.add_run('Backup: Automático e redundante\n')
    specs_cloud.add_run('SLA: 99.95% uptime\n')

    doc.add_paragraph()

    # Cost breakdown
    doc.add_heading('Custos:', 2)

    doc.add_paragraph().add_run('Investimento Inicial:').bold = True
    initial_cloud = doc.add_paragraph(style='List Bullet')
    initial_cloud.add_run('R$ 0 (zero investimento)\n').bold = True

    doc.add_paragraph().add_run('Custos Mensais:').bold = True
    monthly_cloud = doc.add_paragraph(style='List Bullet')
    monthly_cloud.add_run('Servidor 24/7 (4 OCPU): R$ 460/mês\n')
    monthly_cloud.add_run('Processing (~30h/mês): R$ 22/mês\n')
    monthly_cloud.add_run('TOTAL MENSAL: R$ 482/mês\n').bold = True

    doc.add_paragraph().add_run('Total 5 Anos: R$ 29.000').bold = True
    doc.add_paragraph().add_run('Economia vs On-Premise: R$ 24.380 (46%)').bold = True

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    # VANTAGENS
    doc.add_heading('✅ VANTAGENS', 2)
    vant_cloud = doc.add_paragraph(style='List Bullet')
    vant_cloud.add_run('Zero investimento inicial\n')
    vant_cloud.add_run('SLA 99.95% (melhor que on-premise)\n')
    vant_cloud.add_run('Backup automático e redundante\n')
    vant_cloud.add_run('Disaster recovery incluído\n')
    vant_cloud.add_run('Escalabilidade instantânea (aumenta recursos em minutos)\n')
    vant_cloud.add_run('Sem preocupação com manutenção de hardware\n')
    vant_cloud.add_run('Sem custo de energia\n')
    vant_cloud.add_run('Sem risco de falha de hardware\n')
    vant_cloud.add_run('10 TB/mês de transfer GRÁTIS\n')
    vant_cloud.add_run('Datacenter em São Paulo (baixa latência Brasil)\n')
    vant_cloud.add_run('Paga só pelo que usa (backfill on-demand)\n')
    vant_cloud.add_run('Upgrades sem custo adicional\n')
    vant_cloud.add_run('Free Tier para dev/staging\n')
    vant_cloud.add_run('Acesso global de qualquer lugar\n')

    # DESVANTAGENS
    doc.add_heading('❌ DESVANTAGENS', 2)
    desv_cloud = doc.add_paragraph(style='List Bullet')
    desv_cloud.add_run('Custo mensal contínuo (não "paga uma vez")\n')
    desv_cloud.add_run('Dependência de internet estável\n')
    desv_cloud.add_run('Vendor lock-in (migrar provedor é trabalhoso)\n')
    desv_cloud.add_run('Latência de rede (10-50ms vs 0ms local)\n')
    desv_cloud.add_run('Menos controle físico sobre dados\n')
    desv_cloud.add_run('Custos podem aumentar se houver muito egress\n')
    desv_cloud.add_run('Precisa de expertise em cloud (curva de aprendizado)\n')
    desv_cloud.add_run('Suporte técnico por tickets (não presencial)\n')

    add_page_break(doc)

    # ======================
    # NUVEM - HÍBRIDO ULTRA
    # ======================
    doc.add_heading('🔀 Opção 3: Híbrido Ultra-Otimizado', 1)

    # Specs
    doc.add_heading('Configuração:', 2)
    specs_hybrid = doc.add_paragraph(style='List Bullet')
    specs_hybrid.add_run('API Serving: Oracle Free Tier (2 OCPU, 12GB) - GRÁTIS\n')
    specs_hybrid.add_run('Storage ativo: AWS S3 Standard 500 GB\n')
    specs_hybrid.add_run('Storage arquivo: AWS S3 Glacier 2 TB\n')
    specs_hybrid.add_run('Processing: AWS Spot instances (~30h/mês)\n')
    specs_hybrid.add_run('CDN: Cloudflare Free\n')

    doc.add_paragraph()

    # Cost breakdown
    doc.add_heading('Custos:', 2)

    doc.add_paragraph().add_run('Investimento Inicial:').bold = True
    initial_hybrid = doc.add_paragraph(style='List Bullet')
    initial_hybrid.add_run('R$ 0 (zero investimento)\n').bold = True

    doc.add_paragraph().add_run('Custos Mensais:').bold = True
    monthly_hybrid = doc.add_paragraph(style='List Bullet')
    monthly_hybrid.add_run('API Serving (Oracle Free): R$ 0/mês\n')
    monthly_hybrid.add_run('Storage S3: R$ 107/mês\n')
    monthly_hybrid.add_run('Processing (Spot): R$ 36/mês\n')
    monthly_hybrid.add_run('CDN (Cloudflare): R$ 0/mês\n')
    monthly_hybrid.add_run('TOTAL MENSAL: R$ 143/mês\n').bold = True

    doc.add_paragraph().add_run('Total 5 Anos: R$ 8.580').bold = True
    doc.add_paragraph().add_run('Economia vs On-Premise: R$ 44.800 (84%)').bold = True

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    # VANTAGENS
    doc.add_heading('✅ VANTAGENS', 2)
    vant_hybrid = doc.add_paragraph(style='List Bullet')
    vant_hybrid.add_run('Custo MUITO baixo (R$ 143/mês)\n')
    vant_hybrid.add_run('Zero investimento inicial\n')
    vant_hybrid.add_run('Aproveita o melhor de cada provedor\n')
    vant_hybrid.add_run('Serving GRÁTIS (Oracle Free Tier)\n')
    vant_hybrid.add_run('Storage barato (S3 Glacier)\n')
    vant_hybrid.add_run('Processing com desconto (Spot 70% off)\n')
    vant_hybrid.add_run('CDN global grátis (Cloudflare)\n')
    vant_hybrid.add_run('Máxima economia (84% vs on-premise)\n')
    vant_hybrid.add_run('Backup automático e redundante\n')
    vant_hybrid.add_run('Escalabilidade instantânea\n')
    vant_hybrid.add_run('Sem vendor lock-in total (multi-cloud)\n')
    vant_hybrid.add_run('Sem custo de energia ou hardware\n')

    # DESVANTAGENS
    doc.add_heading('❌ DESVANTAGENS', 2)
    desv_hybrid = doc.add_paragraph(style='List Bullet')
    desv_hybrid.add_run('Complexidade de setup (múltiplos provedores)\n')
    desv_hybrid.add_run('Precisa gerenciar 3 contas (Oracle, AWS, Cloudflare)\n')
    desv_hybrid.add_run('Spot instances podem ser interrompidas\n')
    desv_hybrid.add_run('Free Tier Oracle tem limites (2 OCPU)\n')
    desv_hybrid.add_run('Maior curva de aprendizado\n')
    desv_hybrid.add_run('Mais pontos de falha potenciais\n')
    desv_hybrid.add_run('Suporte fragmentado entre provedores\n')
    desv_hybrid.add_run('Latência variável entre serviços\n')

    add_page_break(doc)

    # ======================
    # COMPARAÇÃO LADO A LADO
    # ======================
    doc.add_heading('📊 Comparação Lado a Lado', 1)

    # Create comparison table
    table_comp = doc.add_table(rows=1, cols=4)
    table_comp.style = 'Medium Grid 3 Accent 1'

    # Header
    hdr = table_comp.rows[0].cells
    hdr[0].text = 'Critério'
    hdr[1].text = 'On-Premise'
    hdr[2].text = 'Oracle Cloud'
    hdr[3].text = 'Híbrido'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, '2E75B5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Comparison criteria
    criteria = [
        ('Investimento Inicial', 'R$ 8.800', 'R$ 0', 'R$ 0'),
        ('Custo Mensal', 'R$ 693', 'R$ 482', 'R$ 143'),
        ('Total 5 Anos', 'R$ 53.380', 'R$ 29.000', 'R$ 8.580'),
        ('SLA/Uptime', '~95%', '99.95%', '99.9%'),
        ('Escalabilidade', 'Limitada', 'Alta', 'Muito Alta'),
        ('Backup', 'Manual', 'Automático', 'Automático'),
        ('Manutenção', 'Sua resp.', 'Zero', 'Mínima'),
        ('Controle', 'Total', 'Médio', 'Médio'),
        ('Latência', '0ms', '10-50ms', '20-100ms'),
        ('Complexidade', 'Baixa', 'Média', 'Alta'),
        ('Vendor Lock-in', 'Não', 'Sim', 'Parcial'),
        ('Disaster Recovery', 'Manual', 'Incluído', 'Incluído'),
    ]

    for criterion, onprem, oracle, hybrid in criteria:
        row = table_comp.add_row().cells
        row[0].text = criterion
        row[1].text = onprem
        row[2].text = oracle
        row[3].text = hybrid

        # Highlight best option
        if criterion == 'Investimento Inicial' and onprem != 'R$ 0':
            set_cell_background(row[2], '90EE90')
            set_cell_background(row[3], '90EE90')
        elif criterion == 'Custo Mensal':
            set_cell_background(row[3], '90EE90')
        elif criterion == 'Total 5 Anos':
            set_cell_background(row[3], '90EE90')
        elif criterion == 'Controle':
            set_cell_background(row[1], '90EE90')
        elif criterion == 'Latência':
            set_cell_background(row[1], '90EE90')
        elif criterion == 'Complexidade' and onprem == 'Baixa':
            set_cell_background(row[1], '90EE90')

    add_page_break(doc)

    # ======================
    # RECOMENDAÇÃO FINAL
    # ======================
    doc.add_heading('🎯 Qual Escolher?', 1)

    doc.add_heading('Escolha On-Premise (Comprar Hardware) se:', 2)
    when_onprem = doc.add_paragraph(style='List Bullet')
    when_onprem.add_run('✅ Você já tem o hardware ou pode investir R$ 8.800 agora\n')
    when_onprem.add_run('✅ Quer controle total físico sobre os dados\n')
    when_onprem.add_run('✅ Precisa de latência zero (tudo local)\n')
    when_onprem.add_run('✅ Tem expertise para gerenciar servidores\n')
    when_onprem.add_run('✅ Tem espaço físico e infraestrutura adequada\n')
    when_onprem.add_run('✅ Não quer dependência de internet\n')
    when_onprem.add_run('✅ Prefere pagar mais no início e menos depois\n')

    doc.add_heading('Escolha Oracle Cloud (Backfill) se:', 2)
    when_oracle = doc.add_paragraph(style='List Bullet')
    when_oracle.add_run('✅ Não quer investir R$ 8.800 agora\n')
    when_oracle.add_run('✅ Quer simplicidade (um só provedor)\n')
    when_oracle.add_run('✅ Precisa de SLA 99.95%\n')
    when_oracle.add_run('✅ Quer escalabilidade rápida\n')
    when_oracle.add_run('✅ Não quer gerenciar hardware\n')
    when_oracle.add_run('✅ Quer backup automático\n')
    when_oracle.add_run('✅ Custo médio aceitável (R$ 482/mês)\n')

    doc.add_heading('Escolha Híbrido Ultra-Otimizado se:', 2)
    when_hybrid = doc.add_paragraph(style='List Bullet')
    when_hybrid.add_run('✅ Quer MÁXIMA economia (R$ 143/mês)\n')
    when_hybrid.add_run('✅ Tem experiência com múltiplos provedores\n')
    when_hybrid.add_run('✅ Pode gerenciar setup mais complexo\n')
    when_hybrid.add_run('✅ Quer aproveitar Free Tiers\n')
    when_hybrid.add_run('✅ Não se importa com latência um pouco maior\n')
    when_hybrid.add_run('✅ Quer evitar vendor lock-in total\n')
    when_hybrid.add_run('✅ Economia é prioridade #1\n')

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)
    doc.add_paragraph()

    # Final recommendation box
    rec_box = doc.add_paragraph()
    rec_box.add_run('💡 RECOMENDAÇÃO GERAL:\n\n').bold = True
    rec_box.add_run(
        'Se você não tem o hardware: Comece com Oracle Cloud Backfill (R$ 482/mês).\n'
        'Após 6 meses, quando estiver confortável, migre para Híbrido (R$ 143/mês) '
        'para maximizar economia.\n\n'
    )
    rec_box.add_run(
        'Se você já tem hardware: Mantenha on-premise (R$ 693/mês) e use nuvem '
        'apenas para backup off-site.'
    )

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    add_page_break(doc)

    # ======================
    # CENÁRIOS ESPECÍFICOS
    # ======================
    doc.add_heading('💼 Decisão por Cenário Específico', 1)

    # Cenário 1
    doc.add_heading('Cenário 1: "Não tenho dinheiro agora"', 2)
    cen1 = doc.add_paragraph()
    cen1.add_run('Situação: ').bold = True
    cen1.add_run('Não tenho R$ 8.800 para investir em hardware.\n')
    cen1.add_run('Escolha: ').bold = True
    cen1.add_run('☁️ Oracle Cloud ou Híbrido\n').bold = True
    cen1.add_run('Por quê: ').bold = True
    cen1.add_run('Zero investimento inicial, começa a usar imediatamente.')

    # Cenário 2
    doc.add_heading('Cenário 2: "Quero controle total"', 2)
    cen2 = doc.add_paragraph()
    cen2.add_run('Situação: ').bold = True
    cen2.add_run('Preciso de controle total sobre dados e hardware.\n')
    cen2.add_run('Escolha: ').bold = True
    cen2.add_run('🖥️ On-Premise\n').bold = True
    cen2.add_run('Por quê: ').bold = True
    cen2.add_run('Controle físico 100%, sem dependência de terceiros.')

    # Cenário 3
    doc.add_heading('Cenário 3: "Quero gastar menos possível"', 2)
    cen3 = doc.add_paragraph()
    cen3.add_run('Situação: ').bold = True
    cen3.add_run('Economia é prioridade #1.\n')
    cen3.add_run('Escolha: ').bold = True
    cen3.add_run('🔀 Híbrido Ultra-Otimizado (R$ 143/mês)\n').bold = True
    cen3.add_run('Por quê: ').bold = True
    cen3.add_run('Economia de R$ 44.800 em 5 anos!')

    # Cenário 4
    doc.add_heading('Cenário 4: "Não quero gerenciar hardware"', 2)
    cen4 = doc.add_paragraph()
    cen4.add_run('Situação: ').bold = True
    cen4.add_run('Não tenho tempo/expertise para manutenção.\n')
    cen4.add_run('Escolha: ').bold = True
    cen4.add_run('☁️ Oracle Cloud\n').bold = True
    cen4.add_run('Por quê: ').bold = True
    cen4.add_run('Zero manutenção, SLA garantido, backup automático.')

    # Cenário 5
    doc.add_heading('Cenário 5: "Preciso escalar rápido"', 2)
    cen5 = doc.add_paragraph()
    cen5.add_run('Situação: ').bold = True
    cen5.add_run('Dados podem crescer 10x em 1 ano.\n')
    cen5.add_run('Escolha: ').bold = True
    cen5.add_run('☁️ Oracle Cloud ou Híbrido\n').bold = True
    cen5.add_run('Por quê: ').bold = True
    cen5.add_run('Escalabilidade instantânea, sem precisar comprar hardware novo.')

    # Cenário 6
    doc.add_heading('Cenário 6: "Já tenho o hardware"', 2)
    cen6 = doc.add_paragraph()
    cen6.add_run('Situação: ').bold = True
    cen6.add_run('Já comprei ou tenho o hardware disponível.\n')
    cen6.add_run('Escolha: ').bold = True
    cen6.add_run('🖥️ On-Premise (obviamente!)\n').bold = True
    cen6.add_run('Por quê: ').bold = True
    cen6.add_run('Investimento já feito, custo mensal menor (R$ 693 vs R$ 890).')

    add_page_break(doc)

    # ======================
    # TIMELINE DE CUSTOS
    # ======================
    doc.add_heading('📈 Evolução de Custos ao Longo do Tempo', 1)

    # Table
    table_timeline = doc.add_table(rows=1, cols=6)
    table_timeline.style = 'Light Grid Accent 1'

    # Header
    hdr_time = table_timeline.rows[0].cells
    headers_time = ['Período', 'On-Premise', 'Oracle Cloud', 'Híbrido', 'Melhor', 'Economia']
    for i, h in enumerate(headers_time):
        hdr_time[i].text = h
        hdr_time[i].paragraphs[0].runs[0].bold = True
        set_cell_background(hdr_time[i], '4472C4')
        hdr_time[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Timeline data
    timeline_data = [
        ('Mês 1', 'R$ 9.493', 'R$ 482', 'R$ 143', 'Híbrido', '-R$ 9.350'),
        ('Mês 6', 'R$ 12.958', 'R$ 2.892', 'R$ 858', 'Híbrido', '-R$ 12.100'),
        ('Ano 1', 'R$ 17.116', 'R$ 5.784', 'R$ 1.716', 'Híbrido', '-R$ 15.400'),
        ('Ano 2', 'R$ 25.432', 'R$ 11.568', 'R$ 3.432', 'Híbrido', '-R$ 22.000'),
        ('Ano 3', 'R$ 33.748', 'R$ 17.352', 'R$ 5.148', 'Híbrido', '-R$ 28.600'),
        ('Ano 5', 'R$ 53.380', 'R$ 29.000', 'R$ 8.580', 'Híbrido', '-R$ 44.800'),
    ]

    for period, onprem, oracle, hybrid, best, saving in timeline_data:
        row = table_timeline.add_row().cells
        row[0].text = period
        row[1].text = onprem
        row[2].text = oracle
        row[3].text = hybrid
        row[4].text = best
        row[5].text = saving

        # Highlight best option
        if best == 'Híbrido':
            set_cell_background(row[3], '90EE90')
            set_cell_background(row[5], 'FFD700')

    doc.add_paragraph()
    doc.add_paragraph('📊 Break-even On-Premise: 13 meses (R$ 8.800 ÷ R$ 693)')

    add_page_break(doc)

    # ======================
    # CONCLUSÃO
    # ======================
    doc.add_heading('🎓 Conclusão Final', 1)

    conclusion = doc.add_paragraph()
    conclusion.add_run('Resumo das 3 Opções:\n\n').bold = True

    # Option 1
    conclusion.add_run('🖥️ On-Premise (Hardware): ').bold = True
    conclusion.add_run('Melhor se você quer controle total e já tem/pode investir R$ 8.800. ')
    conclusion.add_run('Custo: R$ 53.380 em 5 anos.\n\n')

    # Option 2
    conclusion.add_run('☁️ Oracle Cloud (Backfill): ').bold = True
    conclusion.add_run('Melhor equilíbrio entre simplicidade e custo. ')
    conclusion.add_run('Zero investimento, R$ 482/mês. ')
    conclusion.add_run('Custo: R$ 29.000 em 5 anos (46% economia).\n\n')

    # Option 3
    conclusion.add_run('🔀 Híbrido Ultra-Otimizado: ').bold = True
    conclusion.add_run('Máxima economia possível! ')
    conclusion.add_run('R$ 143/mês, mas mais complexo. ')
    conclusion.add_run('Custo: R$ 8.580 em 5 anos (84% economia!).\n\n')

    doc.add_paragraph('─' * 80)
    doc.add_paragraph()

    # Final verdict
    verdict = doc.add_paragraph()
    verdict.add_run('🏆 VEREDICTO:\n\n').bold = True
    verdict.add_run(
        '• Se não tem R$ 8.800: Vá de nuvem (Oracle ou Híbrido)\n'
        '• Se quer gastar menos: Híbrido (R$ 143/mês)\n'
        '• Se quer simplicidade: Oracle Cloud (R$ 482/mês)\n'
        '• Se quer controle: On-Premise (mas custa mais a longo prazo)\n\n'
    )

    verdict.add_run('💡 Recomendação pessoal: ').bold = True
    verdict.add_run(
        'Comece com Oracle Cloud por 6 meses, aprenda a usar, '
        'depois migre para Híbrido quando estiver confortável. '
        'Melhor custo-benefício a longo prazo!'
    )

    doc.add_paragraph()
    doc.add_paragraph('─' * 80)
    doc.add_paragraph()

    # Contact
    contact = doc.add_paragraph()
    contact.add_run(f'Documento gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}\n')
    contact.add_run('Sistema: API Geoespacial (FastAPI + GeoServer + Dask)\n')

    return doc

def main():
    """Main function"""
    print("🚀 Gerando comparação simplificada On-Premise vs Nuvem...")

    try:
        doc = create_simple_comparison()

        output_path = '/opt/geospatial_backend/Comparacao_Hardware_vs_Nuvem.docx'
        doc.save(output_path)

        print(f"✅ Documento gerado com sucesso!")
        print(f"📄 Localização: {output_path}")
        print(f"\n📊 Conteúdo:")
        print("  • Resumo de custos (3 opções)")
        print("  • On-Premise: Vantagens e Desvantagens")
        print("  • Oracle Cloud: Vantagens e Desvantagens")
        print("  • Híbrido: Vantagens e Desvantagens")
        print("  • Comparação lado a lado (12 critérios)")
        print("  • Guia de decisão por cenário")
        print("  • Timeline de custos (5 anos)")
        print("  • Conclusão e recomendação final")
        print("\n💰 Destaques:")
        print("  🥇 Híbrido: R$ 143/mês (84% economia)")
        print("  🥈 Oracle: R$ 482/mês (46% economia)")
        print("  🥉 On-Premise: R$ 693/mês (+ R$ 8.800 inicial)")

    except Exception as e:
        print(f"❌ Erro ao gerar documento: {e}")
        raise

if __name__ == '__main__':
    main()
